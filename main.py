import whisper
import torch
import warnings
import librosa
import os
from deepmultilingualpunctuation import PunctuationModel
import spacy
from numba import cuda
from docx import Document
from datetime import datetime

# Фильтрация предупреждений
warnings.filterwarnings("ignore", category=UserWarning, module="transformers.pipelines.token_classification")
warnings.filterwarnings("ignore", category=FutureWarning, module="torch.nn.modules.module")


def save_to_docx(text, sentences, audio_path):
    """Сохраняет результаты в Word документ"""
    doc = Document()

    # Метаданные документа
    doc.add_heading('Результат транскрибации аудио', level=1)
    doc.add_paragraph(f"Файл: транскрибация")
    doc.add_paragraph(f"Дата обработки: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Полный текст
    doc.add_heading('Полный текст:', level=2)
    doc.add_paragraph(text)

    # Разделенные предложения
    doc.add_heading('Разделенные предложения:', level=2)
    for i, sent in enumerate(sentences, 1):
        doc.add_paragraph(f"{i}. {sent}")

    # Сохранение файла
    doc.save(audio_path)
    print(f"\nРезультаты сохранены в файл: {audio_path}")


def main():
    # Проверка файла
    audio_path = "C:/Users/kosty/PycharmProjects/Golos/.venv/input.mp3"
    if not os.path.exists(audio_path):
        print(f"Файл {audio_path} не найден!")
        return

    # Инициализация устройства
    device = "cuda" if torch.cuda.is_available() else "cpu"
    print(f"Используется устройство: {device}")

    if device == "cuda":
        cuda.close()
        torch.cuda.empty_cache()

    try:
        # Загрузка модели Whisper
        model = whisper.load_model("large").to(device)
        print(f"Модель Whisper загружена на {device}")

        # Параметры транскрибации
        params = {
            "language": "ru",
            "fp16": (device == "cuda"),
            "suppress_tokens": [0, 1, 2],
            "verbose": True,
            "beam_size": 5
        }

        # Транскрибация
        print("\nНачинаю обработку аудио...")
        try:
            audio, sr = librosa.load(audio_path, sr=16000)
            print("Аудио загружено через librosa")
            result = model.transcribe(audio, **params)
        except Exception as e:
            print(f"Ошибка librosa: {e}, использую стандартный метод")
            result = model.transcribe(audio_path, **params)

        raw_text = result["text"]
        print("\nИсходный текст транскрибации:")
        print(raw_text)

        # Восстановление пунктуации
        print("\nВосстанавливаю пунктуацию...")
        punct_model = PunctuationModel()
        punctuated_text = punct_model.restore_punctuation(raw_text)
        print("\nТекст с восстановленной пунктуацией:")
        print(punctuated_text)

        # Разделение на предложения
        print("\nРазделяю текст на предложения...")
        try:
            nlp = spacy.load("ru_core_news_sm")
        except OSError:
            print("Модель spaCy для русского языка не найдена. Устанавливаю...")
            from spacy.cli import download
            download("ru_core_news_sm")
            nlp = spacy.load("ru_core_news_sm")

        doc = nlp(punctuated_text)
        print(doc)
        sentences = [sent.text.strip() for sent in doc.sents if sent.text.strip()]
        print(1)
        save_to_docx(punctuated_text, sentences, "C:/Users/kosty/PycharmProjects/Golos/.venv/output.docx")  # явное преобразование в строку
        print(2)
    except Exception as e:
        print(f"\nКритическая ошибка: {e}")
    finally:
        if device == "cuda":
            torch.cuda.empty_cache()


if __name__ == "__main__":
    # Проверка и установка зависимостей
    required = ['python-docx', 'spacy', 'whisper', 'torch', 'librosa', 'deepmultilingualpunctuation']

    import subprocess
    import sys

    for package in required:
        try:
            __import__(package.split('-')[0])
        except ImportError:
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])

    # Запуск основной программы
    main()